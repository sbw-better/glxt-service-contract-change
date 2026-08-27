from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX_PATH = Path(r"D:\projects\glxt-service-contract-change\docs\合同段落变更类型识别服务项目业务接口文档.docx")
TEMP_PATH = DOCX_PATH.with_name(DOCX_PATH.stem + ".layout-fix.tmp.docx")


def remove_heading_page_breaks(document):
    """取消一级标题的强制换页，让章节自然接续，避免上一页只有少量内容。"""
    changed = 0
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Heading 1":
            continue

        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.keep_with_next = True
        changed += 1
    return changed


def create_restarted_numbering(numbering_root, source_num_id):
    """基于现有编号样式创建一个从1开始的新编号实例。"""
    source_numbers = numbering_root.xpath("./w:num[@w:numId='%s']" % source_num_id)
    if not source_numbers:
        raise ValueError("找不到编号定义 numId=%s" % source_num_id)

    existing_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering_root.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    new_num_id = max(existing_ids) + 1

    new_num = deepcopy(source_numbers[0])
    new_num.set(qn("w:numId"), str(new_num_id))

    for old_override in list(new_num.findall(qn("w:lvlOverride"))):
        new_num.remove(old_override)

    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering_root.append(new_num)
    return new_num_id


def set_paragraph_numbering(paragraph, num_id):
    """给段落设置明确编号，避免不同步骤列表互相续号。"""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = paragraph_properties.find(qn("w:numPr"))
    if number_properties is None:
        number_properties = OxmlElement("w:numPr")
        paragraph_properties.append(number_properties)

    level = number_properties.find(qn("w:ilvl"))
    if level is None:
        level = OxmlElement("w:ilvl")
        number_properties.insert(0, level)
    level.set(qn("w:val"), "0")

    number_id = number_properties.find(qn("w:numId"))
    if number_id is None:
        number_id = OxmlElement("w:numId")
        number_properties.append(number_id)
    number_id.set(qn("w:val"), str(num_id))


def restart_each_numbered_block(document):
    """每组连续的步骤列表单独从1编号，目录和各业务流程互不影响。"""
    numbering_root = document.part.numbering_part.element
    current_block_num_id = None
    block_count = 0
    item_count = 0

    for paragraph in document.paragraphs:
        if paragraph.style.name != "List Number":
            current_block_num_id = None
            continue

        if current_block_num_id is None:
            style_num_properties = paragraph.style.element.pPr.find(qn("w:numPr"))
            if style_num_properties is None:
                raise ValueError("List Number样式缺少编号定义")
            style_num_id = style_num_properties.find(qn("w:numId"))
            if style_num_id is None:
                raise ValueError("List Number样式缺少numId")

            current_block_num_id = create_restarted_numbering(
                numbering_root,
                style_num_id.get(qn("w:val")),
            )
            block_count += 1

        set_paragraph_numbering(paragraph, current_block_num_id)
        item_count += 1

    return block_count, item_count


document = Document(str(DOCX_PATH))
heading_count = remove_heading_page_breaks(document)
numbering_blocks, numbered_items = restart_each_numbered_block(document)

document.save(str(TEMP_PATH))
TEMP_PATH.replace(DOCX_PATH)

print(
    "layout fixed: headings=%d, numbering_blocks=%d, numbered_items=%d, output=%s"
    % (heading_count, numbering_blocks, numbered_items, DOCX_PATH)
)
