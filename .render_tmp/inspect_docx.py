from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


docx_path = Path(r"D:\projects\glxt-service-contract-change\docs\合同段落变更类型识别服务项目业务接口文档.docx")
document = Document(str(docx_path))

for index, paragraph in enumerate(document.paragraphs):
    properties = paragraph._p.pPr
    num_id = None
    ilvl = None
    page_break_before = None
    if properties is not None:
        num_pr = properties.find(qn("w:numPr"))
        if num_pr is not None:
            num_id_node = num_pr.find(qn("w:numId"))
            ilvl_node = num_pr.find(qn("w:ilvl"))
            if num_id_node is not None:
                num_id = num_id_node.get(qn("w:val"))
            if ilvl_node is not None:
                ilvl = ilvl_node.get(qn("w:val"))
        page_break_node = properties.find(qn("w:pageBreakBefore"))
        if page_break_node is not None:
            page_break_before = page_break_node.get(qn("w:val"), "true")

    has_page_break_run = any(
        run._r.find(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']") is not None
        for run in paragraph.runs
    )
    if num_id is not None or page_break_before is not None or has_page_break_run or paragraph.style.name.startswith("Heading"):
        print(
            "{0:04d} style={1!r} numId={2!r} ilvl={3!r} pageBefore={4!r} runPageBreak={5} text={6!r}".format(
                index,
                paragraph.style.name,
                num_id,
                ilvl,
                page_break_before,
                has_page_break_run,
                paragraph.text[:100],
            )
        )

print("\nSELECTED PARAGRAPHS")
for index, paragraph in enumerate(document.paragraphs):
    if (40 <= index <= 57) or (143 <= index <= 151) or (180 <= index <= 200):
        style = paragraph.style
        style_num_id = None
        style_ilvl = None
        style_properties = style.element.pPr
        if style_properties is not None:
            style_num_pr = style_properties.find(qn("w:numPr"))
            if style_num_pr is not None:
                style_num_id_node = style_num_pr.find(qn("w:numId"))
                style_ilvl_node = style_num_pr.find(qn("w:ilvl"))
                if style_num_id_node is not None:
                    style_num_id = style_num_id_node.get(qn("w:val"))
                if style_ilvl_node is not None:
                    style_ilvl = style_ilvl_node.get(qn("w:val"))
        print(
            "{0:04d} style={1!r} styleNumId={2!r} styleIlvl={3!r} text={4!r}".format(
                index,
                paragraph.style.name,
                style_num_id,
                style_ilvl,
                paragraph.text[:120],
            )
        )
