import json
import re
from pathlib import Path

from docx import Document


SOURCE_DOCX = Path(r"C:\Users\hyd\Desktop\丹羿精选12号1期私募证券投资基金基金合同（2024-1）-终版.docx")
OUT_DIR = Path(r"D:\projects\glxt-service-contract-change\docs\demo-materials\contract-sample-tests\build")
OUT_JSON = OUT_DIR / "contract_test_dataset.json"


CODE_GROUPS = [
    "20",
    "25",
    "28",
    "43",
    "57",
    "20;25",
    "20;28",
    "25;28",
    "20;25;28",
    "20;25;28;43;57",
    "35;40;45",
    "26;54;76",
    "11;18",
    "32;33;34",
]


def clean_text(text):
    text = re.sub(r"\s+", " ", text or "").strip()
    text = text.replace("\u3000", " ").strip()
    return text


def text_from_table(table):
    values = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        cells = [cell for cell in cells if cell]
        if cells:
            values.append("；".join(cells))
    return values


def should_keep(text):
    if not text:
        return False
    if len(text) < 18 or len(text) > 900:
        return False
    if re.fullmatch(r"[第\d一二三四五六七八九十百、.．（）()\-\s]+", text):
        return False
    if text in {"目录", "基金合同", "风险揭示书"}:
        return False
    if text.count(".") > 8 and len(text) < 80:
        return False
    return any(ch in text for ch in "基金合同管理人托管人投资者份额资产风险费用信息披露变更")


def split_long(text):
    if len(text) <= 360:
        return [text]
    parts = re.split(r"(?<=[。；;])", text)
    result = []
    current = ""
    for part in parts:
        part = clean_text(part)
        if not part:
            continue
        if len(current) + len(part) <= 360:
            current += part
        else:
            if should_keep(current):
                result.append(current)
            current = part
    if should_keep(current):
        result.append(current)
    return result or [text[:360]]


def variant(text, idx):
    replacements = [
        ("基金管理人", "管理人"),
        ("基金托管人", "托管人"),
        ("投资者", "基金投资者"),
        ("本基金", "该基金"),
        ("基金份额", "份额"),
        ("应当", "应"),
        ("可以", "可"),
        ("不得", "不可"),
    ]
    result = text
    for source, target in replacements:
        if source in result and idx % 2 == 0:
            result = result.replace(source, target, 1)
            break
    prefixes = [
        "经本次合同比对确认，",
        "本次修订后，",
        "在更新后的合同文本中，",
        "与原合同相比，",
    ]
    suffixes = [
        "上述内容用于本次变更识别测试。",
        "该段落作为历史相似样本参与投票。",
        "该表述用于验证语义相近但文本不完全一致的场景。",
        "该样本用于补充知识库演示。",
    ]
    if idx % 3 == 0:
        result = prefixes[idx % len(prefixes)] + result
    if idx % 4 == 0:
        result = result + suffixes[idx % len(suffixes)]
    return result[:1900]


def build_dataset():
    doc = Document(str(SOURCE_DOCX))
    raw = []
    for para in doc.paragraphs:
        value = clean_text(para.text)
        if value:
            raw.append(value)
    for table in doc.tables:
        raw.extend(text_from_table(table))

    seen = set()
    base = []
    for item in raw:
        for part in split_long(item):
            part = clean_text(part)
            if should_keep(part) and part not in seen:
                seen.add(part)
                base.append(part)

    if not base:
        raise RuntimeError("No usable contract paragraphs extracted from source docx")

    normal = []
    target_count = 990
    idx = 0
    while len(normal) < target_count:
        text = base[idx % len(base)]
        if idx < len(base):
            paragraph = text
        else:
            paragraph = variant(text, idx)
        codes = CODE_GROUPS[idx % len(CODE_GROUPS)]
        normal.append({
            "rowNo": len(normal) + 2,
            "paragraph": paragraph,
            "codes": codes,
            "source": "合同样例抽取" if idx < len(base) else "合同样例改写",
            "scenario": "正常导入"
        })
        idx += 1

    duplicate_seed = normal[0]["paragraph"]
    conflict_seed = normal[1]["paragraph"]
    separator_seed = normal[2]["paragraph"]
    long_para = (base[0] + " ") * ((2100 // max(1, len(base[0]))) + 2)
    long_para = long_para[:2105]

    prediction_cases = [
        {
            "name": "EXACT_完全相同段落",
            "paragraph": normal[0]["paragraph"],
            "expectedFocus": "命中Hash，matchType=EXACT，不调用模型，score=1。",
        },
        {
            "name": "SEMANTIC_HIGH_轻微改写",
            "paragraph": variant(normal[3]["paragraph"], 12),
            "expectedFocus": "文本不完全相同，适合观察语义召回和HIGH/CANDIDATE。",
        },
        {
            "name": "SEMANTIC_CANDIDATE_同主题不同表述",
            "paragraph": variant(normal[7]["paragraph"], 15),
            "expectedFocus": "适合观察score与similarity的区别。",
        },
        {
            "name": "STRONG_SINGLE_MATCH_FALLBACK_强相似单条兜底",
            "paragraph": variant(normal[9]["paragraph"], 20),
            "expectedFocus": "如果多标签投票没有达标，但第一名相似度较高，可观察CANDIDATE兜底。",
        },
        {
            "name": "NO_RELIABLE_MATCH_无关文本",
            "paragraph": "本段文字描述办公楼停车位、会议室茶歇安排以及设备巡检计划，与基金合同核心条款没有直接关系。",
            "expectedFocus": "预期无可靠历史证据，返回NO_RELIABLE_MATCH或空候选。",
        },
        {
            "name": "ERROR_EMPTY_PARAGRAPH_空段落",
            "paragraph": "",
            "expectedFocus": "请求体校验失败，提示合同段落不能为空。",
        },
        {
            "name": "ERROR_TOO_LONG_超长段落",
            "paragraph": long_para,
            "expectedFocus": "超过2000字符，预期被业务校验拒绝。",
        },
    ]

    payload = {
        "sourceDocx": str(SOURCE_DOCX),
        "baseParagraphCount": len(base),
        "normalRows": normal,
        "duplicateRows": [
            {"paragraph": duplicate_seed, "codes": normal[0]["codes"]},
            {"paragraph": duplicate_seed, "codes": normal[0]["codes"]},
            {"paragraph": duplicate_seed, "codes": normal[0]["codes"]},
            {"paragraph": normal[4]["paragraph"], "codes": normal[4]["codes"]},
            {"paragraph": normal[4]["paragraph"], "codes": normal[4]["codes"]},
        ],
        "conflictRows": [
            {"paragraph": conflict_seed, "codes": normal[1]["codes"]},
            {"paragraph": conflict_seed, "codes": "43;57"},
            {"paragraph": normal[5]["paragraph"], "codes": "20;25"},
            {"paragraph": normal[5]["paragraph"], "codes": "28;43"},
        ],
        "separatorRows": [
            {"paragraph": separator_seed, "codes": "28，20；25;20"},
            {"paragraph": normal[6]["paragraph"], "codes": "57,43；57"},
            {"paragraph": normal[8]["paragraph"], "codes": "26；54，76"},
        ],
        "invalidRows": [
            {"paragraph": "", "codes": "20"},
            {"paragraph": normal[10]["paragraph"], "codes": ""},
            {"paragraph": normal[11]["paragraph"], "codes": "20 25"},
            {"paragraph": normal[12]["paragraph"], "codes": "X" * 65},
        ],
        "longRows": [
            {"paragraph": long_para, "codes": "20;25"}
        ],
        "overLimitRows": [
            {"paragraph": variant(base[i % len(base)], i + 3000), "codes": CODE_GROUPS[i % len(CODE_GROUPS)]}
            for i in range(1001)
        ],
        "predictionCases": prediction_cases,
    }
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "source": str(SOURCE_DOCX),
        "baseParagraphCount": len(base),
        "normalRows": len(normal),
        "output": str(OUT_JSON)
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    build_dataset()
