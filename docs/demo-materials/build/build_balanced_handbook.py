from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

OUT = r"D:\projects\glxt-service-contract-change\docs\demo-materials\合同段落变更类型识别服务-综合汇报讲稿.docx"


def font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, value, size=10, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(value)
    font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)

for name in ["Normal", "Heading 1", "Heading 2"]:
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
doc.styles["Normal"].font.size = Pt(10)
doc.styles["Heading 1"].font.size = Pt(15)
doc.styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
doc.styles["Heading 2"].font.size = Pt(12)
doc.styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("合同段落变更类型识别服务")
font(r, size=22, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("综合汇报讲稿")
font(r, size=14, color=(88, 96, 110))

para(doc, "定位：这版适合业务 + 技术混合听众。讲法上先用业务问题牵引，再自然带出技术实现；不要把技术页讲成参数清单。")
para(doc, "建议节奏：25 到 30 分钟。前 15 分钟讲整体方案和实现思路，后 8 到 10 分钟现场演示，最后留 3 到 5 分钟讨论准确性验证和知识库完善。")

doc.add_heading("一、开场", level=1)
para(doc, "各位好，今天汇报的是合同段落变更类型识别服务。简单说，它要解决的是一个很实际的问题：合同文档比对之后，会出现很多变化段落，业务人员需要判断这些段落对应哪些变更类型。过去这件事主要靠人工经验，而很多类似判断其实在历史材料里已经出现过。")
para(doc, "所以这次第一版的目标，不是做一个完全替代人的系统，而是把历史上确认过的经验沉淀成知识库。新段落来了以后，系统先找相似历史段落，给出候选变更类型和参考依据，最后仍由业务人员确认。")

doc.add_heading("二、逐页讲法", level=1)
slides = [
    ("1. 标题页", "这页先定调：今天既讲业务价值，也讲核心实现，但不会陷入过细的代码细节。重点是让大家知道这套能力为什么可用、怎么用、结果怎么解释。"),
    ("2. 经验难复用", "从业务场景切入。文档比对以后，变化段落不少，每次都从头判断会花时间。系统要做的是把过去确认过的判断经验找回来，让它在新段落识别时发挥作用。"),
    ("3. 第一版闭环", "这里把业务和技术连起来：业务上看，是历史经验复用；技术上看，是把段落转成向量，做相似检索，再根据历史标签投票。第一版先把这条闭环跑通。"),
    ("4. 导入如何变成知识", "历史 Excel 里每一行是一个已确认样本。系统会先规范化文本、算 Hash，Hash 用来识别完全相同的段落；然后调用统一模型网关生成向量，向量用于识别表达相近的段落。"),
    ("5. 预测流程", "新段落进来后，不是一上来就调用模型。系统先看有没有完全相同的历史段落；如果没有，再生成向量，找 Top 10 相似样本，最后根据相似度做加权投票。这样既节省调用，也保留了可解释依据。"),
    ("6. 结果解释", "这里要重点讲清楚两个分数。similarity 是新段落和某条历史段落像不像；score 是相似历史样本对某个类型支持不支持。score 不是概率，它更像历史证据的支持程度。"),
    ("7. 等级和边界", "HIGH 可以作为优先判断方向，CANDIDATE 是候选参考，NO_RELIABLE_MATCH 表示知识库里证据不足。强相似兜底也只是候选，不直接变成高可信结论。"),
    ("8. 为什么不用向量数据库", "技术选择要讲得自然：第一版最多 1 万条历史样本，这个规模下 Java 内存检索足够。先少引入一个组件，可以降低部署、权限和排障复杂度。未来规模扩大再升级。"),
    ("9. 为什么 1024 维、Oracle 和 JVM 索引", "1024 维是效果和成本的平衡；Oracle 用来保存历史样本和审计信息，符合现有系统习惯；JVM 索引让查询更快，导入或重载时整体刷新。"),
    ("10. 模型网关和安全", "模型通过公司统一网关调用，不在本服务里自建模型。API Key 走环境变量，不进代码和演示材料；日志不打印合同正文、向量和敏感信息。模型版本和维度也会记录，避免不同版本混用。"),
    ("11. 知识库完善", "这是新增重点。文档比对后产生的新段落，可以先由系统给候选类型，再由业务确认或修正。只有确认后的段落和类型才补充进知识库，避免错误经验被后续反复引用。"),
    ("12. 现场演示", "演示不要按接口名硬念，而是按业务动作讲：先看知识库状态，再导入样本，再做几类预测，最后看数据库和索引刷新。这样业务和技术听众都能跟上。"),
    ("13. 准确性和上线准备", "最后收口：演示证明流程通了，但效果要靠真实样本验证。建议准备 30 到 100 条未入库段落，先人工标注，再看系统命中、高可信质量和无可靠匹配比例。"),
]
for title, body in slides:
    doc.add_heading(title, level=2)
    para(doc, body)

doc.add_heading("三、技术点的自然表达", level=1)
items = [
    ("Hash", "可以说：它帮助系统快速认出完全相同的段落，避免重复判断和重复入库。"),
    ("Embedding", "可以说：它把段落变成便于比较的语义表示，让文字不完全一样但意思接近的段落也能被找到。"),
    ("Top 10 投票", "可以说：系统不是只看一条历史记录，而是综合最相近的一组历史样本，看哪些类型得到更多支持。"),
    ("1024 维", "可以说：这是在识别效果和资源开销之间做的平衡，适合第一版样本规模。"),
    ("Oracle + JVM 索引", "可以说：Oracle 负责把知识库稳稳保存下来，JVM 索引负责让在线识别更快。"),
]
for k, v in items:
    p = doc.add_paragraph()
    r = p.add_run(k + "：")
    font(r, size=10.5, bold=True)
    p.add_run(v)

doc.add_heading("四、现场演示串词", level=1)
demo = [
    "先看索引状态：这里相当于确认知识库现在是否可用，有多少历史样本参与识别。",
    "导入匿名 Excel：这一步相当于把人工确认过的历史经验放进系统。",
    "做完全一致预测：展示系统可以直接复用历史确认结果。",
    "做相似段落预测：展示系统如何返回候选类型和相似历史依据。",
    "做无关段落预测：展示证据不足时系统不会硬给结论。",
    "看 Oracle 只读查询：给技术同事确认样本确实保存下来，模型版本和向量信息也有记录。",
    "最后讲知识库完善：新段落经过人工确认后，可以再导入，让后续识别更贴近真实业务。",
]
for item in demo:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("五、容易被问到的问题", level=1)
qa = [
    ("为什么不自动入库？", "因为系统输出是建议，不是最终业务结论。合同场景里，只有人工确认后的内容才适合作为后续依据。"),
    ("没有匹配是不是失败？", "不是。它说明当前知识库里没有足够相似的历史样本，反而能提示我们补充知识库。"),
    ("为什么第一版不做得更复杂？", "第一版样本规模和业务目标都比较明确，先用轻量方案把闭环跑通，后续根据真实使用数据再扩展更稳。"),
    ("怎么判断系统效果？", "用未入库的真实段落做人工标注测试，统计命中情况、高可信结果质量和无可靠匹配比例。"),
]
for q, a in qa:
    p = doc.add_paragraph()
    r = p.add_run("问：" + q)
    font(r, size=10.5, bold=True)
    para(doc, "答：" + a)

doc.save(OUT)
print(OUT)
