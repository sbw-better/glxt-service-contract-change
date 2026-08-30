from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

OUT = r"D:\projects\glxt-service-contract-change\docs\demo-materials\合同段落变更类型识别服务-业务汇报讲稿.docx"


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, size=10.5, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
for style_name in ["Normal", "Heading 1", "Heading 2"]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
styles["Heading 2"].font.size = Pt(13)
styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("合同段落变更类型识别服务")
set_font(r, size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("业务汇报讲稿")
set_font(r, size=14, color=(88, 96, 110))

add_para(doc, "使用建议：这份讲稿适合业务和技术混合听众，但主线按业务语言来讲。技术细节只在被问到时补充，不主动展开。")
add_para(doc, "整体节奏：先讲为什么需要，再讲怎么用、怎么看结果，最后讲知识库如何持续完善。")

doc.add_heading("一、开场建议", level=1)
add_para(doc, "各位好，今天主要汇报合同段落变更类型识别服务。这个服务的目标不是替代业务人员做最终判断，而是把过去已经确认过的段落和变更类型沉淀下来。以后遇到相似段落时，系统先给出候选类型和相似历史依据，帮助大家更快完成判断。")
add_para(doc, "所以我今天不会把重点放在技术名词上，而是从业务使用角度说明三件事：它解决什么问题，现场怎么用，结果应该怎么看。最后再补充一下，文档比对后产生的新段落，如何经过人工确认后补充进知识库。")

doc.add_heading("二、逐页自然讲法", level=1)
slides = [
    ("1. 标题页", "这页简单带过。可以说：这次汇报的是业务版，我们重点看这个服务怎么帮助合同段落变更类型识别，而不是展开底层实现细节。"),
    ("2. 历史经验难复用", "从日常工作切入：文档比对完成后，变化段落往往不少。每次都靠人工重新判断，效率会受影响；更可惜的是，很多类似情况过去其实已经判断过，只是没有被方便地复用起来。"),
    ("3. 第一版闭环", "这一页讲服务的基本闭环：先导入历史样本，系统形成知识库；新段落来了以后，系统返回候选变更类型和相似历史段落；人工确认后，新的有效样本还能继续补充进去。"),
    ("4. 业务侧怎么用", "这里讲得朴素一点：业务人员只需要准备确认过的历史段落和类型。识别新段落时，系统会给建议和依据。最终判断仍然由人确认，确认后的新样本可以继续沉淀。"),
    ("5. 结果怎么看", "把结果解释成人能理解的两层：第一，看新段落和历史段落像不像；第二，看历史证据是否支持某些变更类型。HIGH 可以优先关注，CANDIDATE 需要结合上下文再判断。"),
    ("6. 证据不足时不硬给结论", "这页很重要。可以说：系统不是为了每次都给一个看起来确定的答案。如果历史库里没有足够相似的依据，它会提示没有可靠匹配。这反而能减少误导。"),
    ("7. 现场演示顺序", "演示时按业务流程讲，不要按接口名硬念。先看知识库是否有样本，再导入样本，然后试几类段落：完全相同、相似、有候选、无可靠匹配。"),
    ("8. 知识库完善", "这页是本次新增的重点。文档比对完成后，新出现的有价值段落可以整理出来，由业务人员确认变更类型，再补充进知识库。这样系统不是一次性工具，而是会随着业务使用逐步积累。"),
    ("9. 为什么不自动入库", "用业务语言解释：系统建议只是辅助判断，不等于业务确认。如果错误建议直接进入知识库，后面相似段落可能继续参考这个错误。合同场景需要稳，所以必须人工确认后再沉淀。"),
    ("10. 准确性验证", "这里不要承诺没有依据的准确率。可以说：演示能证明流程跑通，但要评价效果，需要拿一批没有入库的真实段落，由业务先标注标准答案，再和系统输出对比。"),
    ("11. 收尾", "最后收在三个点：能复用历史经验，能给新段落提供候选类型和参考依据，能通过人工确认后的新样本持续完善。第一版适合先试用，再用真实样本验证效果。"),
]
for title, body in slides:
    doc.add_heading(title, level=2)
    add_para(doc, body)

doc.add_heading("三、现场演示时可以这样串", level=1)
demo = [
    "第一步，先看知识库状态。这里主要看当前有没有历史样本可用，不需要展开技术字段。",
    "第二步，导入一份匿名历史样本。可以说明：这就相当于把过去人工确认过的经验放进知识库。",
    "第三步，拿一段完全相同的文本做识别。系统直接返回历史类型，说明重复场景可以快速复用。",
    "第四步，拿一段表达相近但文字不同的文本做识别。重点看系统返回的候选类型和相似历史段落。",
    "第五步，拿一段明显无关的文本做识别。这里要强调：没有足够依据时，系统不会硬给结论。",
    "最后补一句：文档比对产生的新段落，后续也可以按这个方式整理、确认、导入，让知识库持续变完整。",
]
for item in demo:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("四、知识库完善的推荐说法", level=1)
add_para(doc, "文档比对不是流程终点。比对完成后，业务人员会看到一些新的、有代表性的变化段落。这些段落如果经过人工确认了变更类型，就可以作为新的历史样本补充进知识库。")
add_para(doc, "这里我们坚持一个原则：系统预测结果不直接自动入库。因为预测只是建议，只有人工确认过的内容，才适合作为后续识别的依据。这样做看起来多了一步，但能保证知识库质量。")
add_para(doc, "后续使用得越多，确认过的样本越丰富，系统给出的参考依据也会越来越贴近真实业务场景。")

doc.add_heading("五、准确性验证的推荐说法", level=1)
add_para(doc, "目前还没有真实人工确认测试集，所以汇报时不要说准确率已经达到多少。更稳妥的说法是：演示样本证明流程可用，正式效果需要用真实业务样本验证。")
add_para(doc, "建议准备 30 到 100 条没有进入历史库的段落，由业务人员先标注标准变更类型，再看系统是否命中、哪些结果是高可信、哪些提示无可靠匹配。这样得到的结论才适合用于验收和后续优化。")

doc.add_heading("六、被问到技术问题时的简短回答", level=1)
brief = [
    ("为什么能找相似段落？", "系统会把段落转成便于比较的语义表示，再从历史知识库里找相近内容。"),
    ("为什么结果还要人工确认？", "因为合同变更类型最终是业务判断，系统提供的是候选建议和参考证据。"),
    ("为什么不自动学习？", "自动学习容易把错误判断沉淀进去。人工确认后再入库，更稳，也更适合合同场景。"),
    ("没有匹配结果是不是系统失败？", "不是。它说明当前知识库里没有足够可靠的历史依据，反而提醒我们可能需要补充样本。"),
]
for q, a in brief:
    p = doc.add_paragraph()
    r = p.add_run("问：" + q)
    set_font(r, size=10.5, bold=True)
    add_para(doc, "答：" + a)

doc.save(OUT)
print(OUT)
