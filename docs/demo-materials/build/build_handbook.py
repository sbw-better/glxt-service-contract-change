from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"D:\projects\glxt-service-contract-change\docs\demo-materials\合同段落变更类型识别服务-汇报讲稿与演示手册.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.name = "Calibri"
styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
styles["Heading 2"].font.name = "Calibri"
styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
styles["Heading 2"].font.size = Pt(13)
styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("合同段落变更类型识别服务")
r.bold = True
r.font.size = Pt(22)
r.font.name = "Calibri"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("汇报讲稿与线上演示手册")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(91, 100, 114)
r.font.name = "Calibri"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

doc.add_paragraph("适用场景：业务 + 技术混合听众，线上会议投屏，预计 25-30 分钟。")
doc.add_paragraph("材料口径：本轮只准备演示材料，不修改业务代码；当前已知编译 P0 问题必须在正式演示前修复并保存验证证据。")

doc.add_heading("一、建议时间分配", level=1)
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(["环节", "建议时长", "目标", "备注"]):
    set_cell_text(hdr[i], text, True)
    set_cell_shading(hdr[i], "E8EEF5")
rows = [
    ["开场与背景", "3 分钟", "说明为什么需要该能力", "从人工识别慢、历史经验难复用切入"],
    ["方案与边界", "8 分钟", "讲清楚系统做什么、不做什么", "强调第一版简单可用"],
    ["算法与解释", "6 分钟", "解释 similarity、score、HIGH/CANDIDATE", "避免把 score 说成概率"],
    ["现场演示", "8-10 分钟", "展示可用流程和可观测证据", "Swagger/Postman + Oracle 只读查询"],
    ["风险与下一步", "3 分钟", "讲清楚编译修复、真实测试集和知识库完善", "用事实建立可信度"],
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        set_cell_text(cells[i], text)

doc.add_heading("二、逐页讲稿", level=1)
slides = [
    ("1. 标题页", "本次汇报的是合同段落变更类型识别服务第一版。它不是替代人工审查，而是把历史段落和变更类型经验沉淀下来，在新段落出现时给出候选类型和参考证据。"),
    ("2. 背景价值", "合同文档比对完成后会产生大量变化段落。过去依赖人工逐段判断，耗时且口径不稳定。系统的目标是复用历史经验，提升初筛效率。"),
    ("3. 第一版边界", "第一版刻意保持轻量：历史 Excel 导入、统一模型网关向量化、Oracle 持久化、JVM 内存检索。不引入向量数据库、Rerank、异步任务和额外业务表。"),
    ("4. 核心闭环", "历史样本先进入知识库，新段落预测时匹配历史样本。文档比对产生的新段落，只有经过人工确认类型后才补充回历史库，这样系统能持续完善，同时避免误判自我放大。"),
    ("5. 接口结构", "导入、预测、索引状态和索引重载四类接口覆盖第一版核心使用场景。导入和预测必须携带 UserId，用于模型平台审计透传，不入库也不输出日志。"),
    ("6. 预测算法", "系统先做规范化和 Hash 精确匹配；未命中才调用模型生成向量。语义匹配时取 Top 10 历史样本，过滤低相似度后做 similarity 平方加权投票。"),
    ("7. 分数解释", "similarity 是新段落和某条历史段落有多像；score 是历史证据对某个类型的综合支持度。score 不是概率，多标签分数之和不要求等于 1。"),
    ("8. 技术选型", "当前样本规模最多 1 万条，JVM 内存精确点积足够。1024 维能把网络、存储、内存和计算压力控制在合理范围。Oracle 保留业务可审计数据。"),
    ("9. 模型与安全", "模型通过公司统一网关接入。API Key 从环境变量读取，不进入代码、Git 或日志。对异常重试也做了控制，避免无意义重试和敏感信息泄露。"),
    ("10. 知识库完善", "这一页回应后续演进：文档比对后的候选段落可以导出为待确认样本，人工确认后再入库。近期可复用 Excel 导入，后续再考虑确认接口和审计字段。"),
    ("11. 演示顺序", "现场按可用性、证据、边界推进。先看索引状态，再导入样本，再查 Oracle，再分别演示 EXACT、HIGH、CANDIDATE、NO_RELIABLE_MATCH 和 reload。"),
    ("12. 演示前风险", "这里要如实说明：目前 Java 8 下 mvn test 编译失败，主因是缺少 util 包 import。正式演示前必须修复、测试和打包成功。"),
    ("13. 准确性验证", "目前没有真实测试数据，因此不能宣称准确率。建议准备 30-100 条人工标注段落，统计覆盖率、准确率、召回率、F1、HIGH 准确率和无可靠匹配比例。"),
    ("14. 收尾结论", "第一版适合先上线试用，用人工确认样本持续完善知识库。上线前关键工作是修复编译、准备匿名样本、完成独立验证表和备用录屏。"),
]
for title, script in slides:
    doc.add_heading(title, level=2)
    doc.add_paragraph(script)

doc.add_heading("三、现场演示操作脚本", level=1)
steps = [
    ("准备窗口", "打开 PPT、Swagger 或 Postman、Oracle 查询窗口、服务日志窗口。确认不要展示 API Key、真实合同正文和真实 UserId。"),
    ("index/status", "调用 GET /glxt-service-contract-change/service/contract-change/index/status。说明 status、sampleCount、modelVersion、vectorDimension、loadedAt、errorCount。"),
    ("导入匿名小 Excel", "调用 POST /samples/import，文件使用正常样本 Excel，Header 带 UserId。观察 totalRows、inserted、updated、skipped、indexReloaded。"),
    ("Oracle 只读查询", "查询 TEXT_HASH、BGLX_CODES、VECTOR_DIM、MODEL_VERSION、DBMS_LOB.GETLENGTH(VECTOR_DATA)。说明 1024 维 Float32 小端序应约 4096 字节。"),
    ("EXACT 预测", "用已导入段落原文预测。预期 matchType=EXACT，score=1，similarity=1，不调用模型。"),
    ("HIGH 语义预测", "用语义相近但文字不同的段落预测。解释相似历史段落、supportCount 和 HIGH 条件。"),
    ("CANDIDATE/兜底", "用强相似但标签分散的段落预测。解释第一名 similarity 足够高时触发候选兜底，score 仍是真实投票占比。"),
    ("NO_RELIABLE_MATCH", "用明显不相关段落预测。说明没有足够可靠历史证据时系统不强行给结论。"),
    ("index/reload", "调用 POST /index/reload。说明用于模型版本切换、手工恢复或数据库样本变更后的索引刷新。"),
]
for title, body in steps:
    p = doc.add_paragraph(style=None)
    r = p.add_run(title + "：")
    r.bold = True
    p.add_run(body)

doc.add_heading("四、常见追问回答", level=1)
qa = [
    ("为什么不用向量数据库？", "第一版样本上限 1 万条，JVM 内存精确检索足够。少一个组件，部署、备份、权限和故障定位都更简单。"),
    ("为什么类型编码不做向量？", "类型编码不是自然语言证据，本项目要判断的是段落和历史段落是否相似，再由历史段落标签投票。"),
    ("为什么不自动把预测结果入库？", "预测结果未经人工确认会污染历史标签。系统支持人工确认后入库，保证知识库越用越准，而不是越用越偏。"),
    ("没有真实测试数据怎么证明效果？", "现场演示证明流程可用；准确性要通过 30-100 条独立人工标注段落统计指标后再给结论。"),
    ("如果模型版本变化怎么办？", "MODEL_VERSION 和 VECTOR_DIM 用于隔离向量。模型或维度变化后需要重新生成历史向量，不能混合检索。"),
    ("强相似兜底为什么还是 CANDIDATE？", "因为它只是说明第一条历史样本很像，不代表多条历史证据形成稳定共识，所以不能直接给 HIGH。"),
]
for q, a in qa:
    p = doc.add_paragraph()
    r = p.add_run("问：" + q)
    r.bold = True
    doc.add_paragraph("答：" + a)

doc.add_heading("五、演示前硬性检查", level=1)
checks = [
    "修复当前编译 P0：缺少 ContractTextNormalizer、HashUtils、ChangeTypeCodes、VectorCodec、VectorUtils 等 import。",
    "使用 Java 8 执行 mvn clean test，并保存成功截图或日志摘要。",
    "执行 mvn -DskipTests package，并确认 target jar 能启动。",
    "确认 Oracle 测试库可访问，TPIF_HTDLYB 表存在，当前账号只读查询可用。",
    "确认 Embedding 网关、EMBEDDING_API_KEY、UserId 请求头可用。",
    "准备匿名演示 Excel，避免展示真实合同、真实人员、真实密钥。",
    "准备备用录屏和静态响应截图，防止线上网络或模型网关临时不可用。",
]
for item in checks:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("六、人工确认入库口径", level=1)
doc.add_paragraph("文档比对完成后，系统可以把候选变更段落导出为待确认样本。业务人员确认或修正变更类型后，再通过现有 Excel 导入流程沉淀到 TPIF_HTDLYB。")
doc.add_paragraph("入库条件建议：段落有效、类型合法、Hash 去重、冲突人工处理、记录模型版本和维度。未经确认的预测结果不自动入库。")

doc.save(OUT)
print(OUT)
